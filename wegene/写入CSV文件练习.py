# -*- coding: utf-8 -*-
# @Time    : 2020/6/17 15:35
# @Author  : qinshuaibo
"""
import requests
from  bs4 import BeautifulSoup
from browsermobproxy import Server
from selenium import webdriver
import time
from selenium.webdriver.chrome.options import Options
driver = webdriver.Chrome()
driver2 = webdriver.Chrome()
driver.get('https://www.wegene.com/demo2/male/report/%E5%81%A5%E5%BA%B7%E9%A3%8E%E9%99%A9')
# driver.get('https://www.wegene.com/demo2/male/report/%E5%81%A5%E5%BA%B7%E9%A3%8E%E9%99%A9')
time.sleep(1)
#获取完整渲染的网页源代码
pageSource = driver.page_source
#打印pageSource的类型
print(type(pageSource))
#打印pageSource
# print(pageSource)
#疾病名称
titles = driver.find_elements_by_xpath("//div[@class='title']")
for u in titles:
    print(u.text)
#超链接
#获取所有a标签
links = driver.find_elements_by_xpath("//div[2]/div/div/div[2]/div[3]/div/ul/li[2]/a")
#列表长度，一共有多少个a标签
length = len(links)
#遍历列表循环，使程序可以逐一点击
for i in range(0, length):
    #每次循环内部都重新获取a标签
    links = driver.find_elements_by_xpath("//div[2]/div/div/div[2]/div[3]/div/ul/li[2]/a")
    #把列表中的a标签赋给link
    link = links[i]
    #提取a标签中的链接，这里提取到的链接是文本无法点击
    url = link.get_attribute('href')
    #在这不能使用click，因为没用，直接用浏览器打开
    driver.get(url)
    #留出加载时间
    time.sleep(1)
    #用text的意思是输出纯文本
    titles = driver.find_element_by_xpath("//div[@class='title']").text

    content = driver.find_element_by_xpath("//span[@class='color-wegene-blue font-weight-bold']")
    print (url.get_attribute("href"))
#关闭浏览器
driver.close()
"""

# -*- coding: utf-8 -*-
# @Time    : 2020/7/14 20:54
# @Author  : qinshuaibo
# -*- coding: utf-8 -*-
# @Time    : 2020/7/14 18:55
# @Author  : qinshuaibo
import requests
from bs4 import BeautifulSoup
from browsermobproxy import Server
from selenium import webdriver
from selenium.webdriver.chrome.options import Options
import time
import csv
import codecs
# 创建并写入word文档
import docx

# 创建内存中的word文档对象
file = docx.Document()
# 实例化Options对象
chrome_options = Options()
# 把Chrome浏览器设置为静默模式
chrome_options.add_argument('--headless')
driver = webdriver.Chrome(options = chrome_options)
driver2 = webdriver.Chrome()
driver.get('https://www.abstractsonline.com/pp8/#!/9045/session/862')
time.sleep(5)
# 获取完整渲染的网页源代码
pageSource = driver.page_source
# 打印pageSource的类型
print(type(pageSource))
# 打印pageSource
# print(pageSource)
# 集合名称
# titles = driver.find_elements_by_xpath("//div[@class='title']")
# titles = driver.find_elements_by_xpath("//body/div[7]/div[2]/div/ul[1]/li/a")
titles = driver.find_elements_by_xpath("//div/div[2]/div[2]/table/tbody/tr/td[2]/a")

titles_shuliang = len(titles)
print(titles_shuliang)
for title in titles:

    if len(title.get_attribute("innerText")) > 20:
        print(title.get_attribute("innerText"))
        print(len(title.get_attribute("innerText")))

# 超链接
# 获取所有带链接的标签
links = driver.find_elements_by_xpath("//div/div[2]/div[2]/table/tbody/tr/td[2]/a")
# 列表长度，一共有多少个a标签
length = len(links)
print(length)
print("-" * 50)


def create_docx(name_list, content_list):
    document = Document()
    paragraph = document.add_


for link in links:
    url = link.get_attribute('href')

    # print(url)
    driver2.get(url)
    time.sleep(3)
    names = driver2.find_elements_by_xpath("//div/div[1]/div/div[3]/div[1]/h1")
    length_names = len(names)

    for name in names:
        if len(name.get_attribute("innerText")) > 20:
            aa = name.get_attribute("innerText")
            print(name.get_attribute("innerText"))
            contents = driver2.find_elements_by_xpath("//div/div[2]/div[2]/dl/dd[2]")
            for content in contents:
                print(content.get_attribute("innerText"))

                with open('test.docx', 'w', encoding = 'utf-8'):
                    file.write('\n'.join([aa]))
                # # 写入若干段落
                # file.add_paragraph(content.get_attribute("innerText"))
                # # 保存
                # file.save("D:\aa.docx")

                print('\n')

# 关闭浏览器
driver.close()
driver2.close()
